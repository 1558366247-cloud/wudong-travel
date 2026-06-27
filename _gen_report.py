# -*- coding: utf-8 -*-
"""生成课程设计大作业报告"""
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import datetime

doc = Document()

# ===== 设置默认字体 =====
style = doc.styles['Normal']
font = style.font
font.name = '宋体'
font.size = Pt(12)
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.first_line_indent = Pt(24)

# ===== 封面 =====
for _ in range(4):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('计算机工程学院/大数据学院')
run.font.size = Pt(16)
run.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('《人工智能工程实践》')
run.font.size = Pt(18)
run.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('课程设计（大作业）报告')
run.font.size = Pt(18)
run.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('2025 - 2026 学年第 2 学期')
run.font.size = Pt(14)

doc.add_paragraph()

info_items = [
    ('项目名称', '乌东文旅"衣食住行"综合服务平台智能开发'),
    ('项目组号', '第4组'),
    ('专    业', '计算机科学与技术'),
    ('班    级', '2023计算机科学与技术1班'),
    ('姓    名', '（待填写）'),
    ('学    号', '（待填写）'),
    ('班级序号', '（待填写）'),
    ('指导教师', '王佛琴'),
    ('日    期', '2026年7月5日'),
]

for label, value in info_items:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Pt(0)
    run = p.add_run(f'{label}：{value}')
    run.font.size = Pt(14)

doc.add_page_break()

# ===== 课程设计目的 =====
doc.add_heading('课程设计（大作业）目的', level=1)
p = doc.add_paragraph()
p.add_run(
    '本课程旨在实现从基础技能到高阶素养的递进式培养目标。帮助学生建立正确的AI协作观，'
    '使其能熟练借助AI提升编码效率，同时具备审查、重构和优化AI生成代码的工程判断力，'
    '而非盲目依赖工具；深化学生对软件工程本质的理解，打破"编码即开发"的误区，'
    '强化其在需求理解、架构设计、安全风控及团队协作方面的核心能力；致力于塑造学生的'
    '系统性思维，使其跳出纯技术视角，懂得如何通过技术方案优化游客体验与商家运营效率，'
    '并能基于模拟运营数据进行商业闭环分析，最终实现从"技术执行者"到"业务驱动者"的认知跃升。'
)

doc.add_heading('课程设计（大作业）具体要求', level=1)
p = doc.add_paragraph()
p.add_run(
    '本课程以"乌东文旅衣食住行综合服务平台"的真实企业需求为核心载体，重点训练学生运用'
    'AI辅助编程工具（如大模型代码生成、智能测试用例编写等）贯穿环境搭建、业务编码等'
    '全流程实操；严格遵循软件开发生命周期（SDLC），涵盖需求分析、系统架构设计、'
    '接口规范定义、代码审查及部署交付等标准化环节；深度剖析文旅行业的信息整合逻辑，'
    '涉及商家入驻、商品管理、订单支付及核销等真实业务节点；通过模拟数据填充与多角色'
    '（游客、商家、平台方）场景演练，让学生体验流量分析、订单处理及运营决策的全过程。'
)

p = doc.add_paragraph()
p.add_run(
    '选题：乌东文旅"衣食住行"综合服务平台分为6个模块，共6个小组，每个小组至少完成一个模块的开发。'
    '本组（第4组）负责"行——线路订票"模块的四端开发，涵盖景区门票、路线套餐、电子票核销、'
    '交通攻略等核心功能模块。'
)

p = doc.add_paragraph()
p.add_run(
    '考核要求：（1）平时成绩40%：其中创建知识库、AI编程工具的安装使用及编码等过程性考核占50%，'
    '答辩占50%。（2）期末成绩60%：项目作品完成情况占60%、个人项目实践报告占40%。'
)

doc.add_page_break()

# ===== 正文 =====
doc.add_heading('课程设计（大作业）具体内容', level=1)

# --- 3.1 ---
doc.add_heading('3.1 引言与项目背景', level=2)

p = doc.add_paragraph()
p.add_run(
    '随着数字经济的快速发展，文旅行业的数字化转型已成为提升游客体验、优化运营效率的'
    '关键路径。贵州乌东村作为苗族文化特色村寨，拥有苗族银饰、蜡染刺绣、苗家长桌宴、'
    '特色民宿、苗寨梯田与节庆文化等丰富的文旅资源，但目前运营以线下为主，缺乏统一的'
    '线上服务入口。本课程设计以校企合作为背景，以乌东文旅"衣食住行"综合服务平台为载体，'
    '旨在构建一站式数字化文旅服务平台，覆盖景区门票、路线套餐、交通攻略等核心业务。'
)

p = doc.add_paragraph()
p.add_run(
    '本项目采用AI辅助开发模式，利用Claude Code等AI编程工具进行需求分析、架构设计、'
    '代码生成、测试用例编写及文档生成的全流程智能化开发。通过AI与人类的协同工作，'
    '在保证代码质量的前提下显著提升开发效率。项目定位为面向游客、商家、平台管理员'
    '三端用户的综合服务系统，核心价值在于以苗族文化为特色，打通"种草-决策-预订-支付-核销-评价"'
    '全链路，为乌东村文旅产业提供数字化基础设施。'
)

p = doc.add_paragraph()
p.add_run(
    '在本项目中，本人担任全栈开发角色，负责第4组"行——线路订票"模块的四端开发：'
    '后端API接口（基于Midway.js框架和TypeORM）、PC Web前端（基于React 18和Ant Design 5）、'
    '微信小程序端（微信原生框架）以及后台管理端（React + Ant Design）。'
    '本模块的业务范围包括：景区门票购买（票种选择→日期选择→数量→游客信息→支付→电子票）、'
    '苗寨一日游/两日游/多日游路线套餐（行程浏览→出发日期→人数→游客信息→支付）、'
    '电子票生成与二维码核销、交通攻略展示等核心功能。通过对企业需求文档的深入理解，'
    '本人对文旅行业的票务系统设计有了全面认识，包括票种管理、库存控制、有效期管理、'
    '实名制购票、二维码核销等完整业务环节。'
)

# --- 3.2 ---
doc.add_heading('3.2 需求分析与可行性研究', level=2)

p = doc.add_paragraph()
p.add_run('3.2.1 业务需求分析')
p = doc.add_paragraph()
p.add_run(
    '本模块面向三类用户角色：游客（注册用户）、商家（景区/线路运营商）和平台管理员。'
    '游客通过小程序端或PC端浏览景区和路线套餐，选择票种或路线后进行实名制购票、'
    '在线支付（微信支付）、获取电子票二维码，到景区凭二维码核销入园。'
    '商家通过管理后台管理景区信息、票种库存与价格、路线套餐与行程安排，以及扫码核销电子票。'
    '平台管理员通过后台管理系统进行全模块数据监控、订单管理与退款审核。'
)

p = doc.add_paragraph()
p.add_run(
    '根据《乌东文旅衣食住行需求规格说明书》，本模块的核心业务规则（R9系列）包括：'
    'R9-01 门票有效期为购买日起指定天数内（由票种配置），超期自动失效不可退；'
    'R9-02 每张电子票对应唯一二维码，核销后状态变为已使用，不可重复使用；'
    'R9-03 路线套餐须提前1天购买；出发日期3天前可免费退；出发前3天内不可退；'
    'R9-04 路线套餐购买时须填写全部游客姓名与身份证号（用于景区实名制入园）；'
    'R9-05 电子票核销须在景区营业时间内进行；核销成功后发送微信消息通知用户。'
    '这些业务规则是系统设计的核心约束，直接影响订单状态机和退款逻辑的设计。'
)

p = doc.add_paragraph()
p.add_run('3.2.2 技术选型与可行性分析')
p = doc.add_paragraph()
p.add_run(
    '后端技术选型：采用Midway.js v3框架（基于Node.js的企业级全栈框架），搭配TypeORM进行'
    '数据库操作，MySQL 8.0作为关系型数据库。Midway.js提供了依赖注入（IoC）、装饰器路由、'
    'Swagger文档自动生成、JWT认证、参数校验等企业级特性，非常适合团队并行开发场景。'
    '前端Web端：React 18 + Ant Design 5 + React Router 6，利用Ant Design的组件体系快速'
    '构建界面，设计规范中的苗银蓝(#1F5FA8)和苗绣橙(#E85D2F)作为品牌色贯穿全局。'
    '微信小程序端：采用微信原生框架（WXML + WXSS + JS），通过wx.request调用后端API。'
    '管理后台：同样基于React + Ant Design，采用侧边栏+内容区的经典管理后台布局。'
)

p = doc.add_paragraph()
p.add_run(
    'AI介入点分析：在本项目中，AI工具（Claude Code）在以下环节发挥了重要作用：'
    '（1）需求文档解读：快速从需求规格说明书中提取模块功能、业务规则及数据实体定义，'
    '生成结构化的需求摘要；（2）数据库设计：根据需求文档中的数据实体描述，自动生成'
    'TypeORM实体类，包含字段定义、关联关系、索引等；（3）API接口生成：根据业务逻辑'
    '自动生成控制器（Controller）和服务层（Service）代码，包含CRUD操作和业务校验；'
    '（4）前端页面生成：根据页面功能描述生成React组件和微信小程序页面；'
    '（5）种子数据生成：模拟乌东地区真实的景区、票种、路线数据。'
)

p = doc.add_paragraph()
p.add_run('3.2.3 功能用例图')
p = doc.add_paragraph()
p.add_run(
    '本模块主要用例包括：游客UC-01浏览景区列表与详情、UC-02浏览路线套餐列表与详情、'
    'UC-03购买门票（选择票种→日期→数量→填写游客信息→支付→生成电子票）、'
    'UC-04预订路线套餐（选择路线→出发日期→人数→填写游客信息→支付）、'
    'UC-05查看电子票二维码（含票号、有效期、使用状态）、UC-06申请退票/退款、'
    'UC-07评价景区/路线（文字+图片+1-5星评分）、UC-08收藏/取消收藏景区/路线/攻略、'
    'UC-09浏览交通攻略（按出发地筛选）、UC-10查看订单列表与详情；'
    '商家UC-11核销电子票（扫码或手动输入票号）、UC-12管理景区/票种/路线信息；'
    '管理员UC-13审核退款申请、UC-14查看数据看板（今日订单/营收/热销排行）。'
)

# --- 3.3 ---
doc.add_heading('3.3 系统设计与架构搭建', level=2)

p = doc.add_paragraph()
p.add_run('3.3.1 业务架构设计')
p = doc.add_paragraph()
p.add_run(
    '本系统采用前后端分离的四层架构：展示层、接口层、业务逻辑层和数据层。'
    '展示层包括PC Web前端（React应用）、微信小程序端和后台管理端（React应用），'
    '三端通过统一的RESTful API调用后端服务。'
    '接口层基于Midway.js的Koa中间件，负责JWT身份认证、CORS跨域处理、请求参数校验'
    '和Swagger文档生成。为了保证接口安全性，敏感操作（下单、退款、核销）需验证JWT Token，'
    '并对敏感接口实施IP级别限流（每分钟每IP不超过10次）。'
    '业务逻辑层包含8个核心服务：ScenicSpotService（景区服务）、TicketTypeService（票种服务）、'
    'RoutePackageService（路线服务）、TravelOrderService（订单服务，含支付回调和库存扣减）、'
    'ETicketService（电子票服务）、ReviewService（评价服务）、FavoriteService（收藏服务）'
    '和TransportGuideService（交通攻略服务）。每个服务独立封装业务逻辑，通过依赖注入进行协作。'
    '数据层使用MySQL数据库，通过TypeORM进行对象关系映射，支持Migrations数据库迁移。'
)

p = doc.add_paragraph()
p.add_run('3.3.2 数据库ER图与核心表结构')
p = doc.add_paragraph()
p.add_run(
    '本模块共设计9张数据表，以景区(t_scenic_spot)和路线套餐(t_route_package)为两个核心聚合根。'
    '景区表包含：scenic_id(主键)、name(景区名称)、address(地址)、longitude/latitude(坐标)、'
    'open_time(开放时间)、description(介绍)、main_image(主图)、status(上下架状态)等字段。'
    '票种表(t_ticket_type)通过scenic_id外键关联景区，包含：ticket_type_id(主键)、name(票种名)、'
    'price(价格)、market_price(市场价)、stock(库存)、daily_stock(日库存上限)、'
    'validity_days(有效期天数)、usage_note(使用须知)等字段。'
    '路线套餐表(t_route_package)包含：route_id(主键)、title(标题)、duration(行程天数)、'
    'themes(主题)、price(价格)、departure/destination(出发地/目的地)、'
    'accommodation_standard(住宿标准)、meal_standard(餐饮标准)、stock(库存)、'
    'sold_count(已售)、rating(评分)等字段。'
    '行程表(t_route_itinerary)通过route_id外键关联路线，包含每天的景点、用餐、住宿、交通安排。'
    '订单表(t_travel_order)以order_no为主键（业务主键），包含订单类型(ticket/route)、'
    '商品信息、金额、游客信息(JSON)、订单状态（pending_pay/paid/refunding/refunded/'
    'completed/cancelled）、支付信息等字段。'
    '电子票表(t_e_ticket)通过order_no和ticket_type_id/route_id关联，包含唯一二维码(qr_code)、'
    '有效期、游客信息、状态（unused/used/refunded/expired）、核销信息等字段。'
    '评价表(t_review)和收藏表(t_favorite)分别记录用户的评价内容和收藏行为。'
)

p = doc.add_paragraph()
p.add_run('3.3.3 接口文档生成')
p = doc.add_paragraph()
p.add_run(
    '利用Midway.js的Swagger集成（@midwayjs/swagger），通过控制器方法上的@ApiTags、'
    '@ApiOperation装饰器自动生成Swagger接口文档。AI工具根据需求文档中的功能描述，'
    '自动为每个接口添加了中文注释和参数说明。后端启动后，可通过访问'
    'http://localhost:7001/swagger-ui/index.html 查看完整的接口文档，'
    '包含所有12组接口的请求参数、响应格式和示例数据。'
)

p = doc.add_paragraph()
p.add_run('3.3.4 部署架构')
p = doc.add_paragraph()
p.add_run(
    '系统部署采用以下架构：Nginx作为反向代理和静态资源服务器，将/api/路径的请求转发到'
    'Midway.js后端服务（运行在7001端口），将静态文件请求直接返回React构建产物。'
    'MySQL数据库和Redis缓存部署在同一内网环境，通过环境变量配置连接信息。'
    '微信小程序通过HTTPS协议访问Nginx的公网地址。整个系统可部署在单台云服务器上，'
    '也可根据负载情况扩展为前后端分离部署。'
)

# --- 3.4 ---
doc.add_heading('3.4 AI辅助开发与核心编码实现', level=2)

p = doc.add_paragraph()
p.add_run('3.4.1 环境搭建与AI初始化配置')
p = doc.add_paragraph()
p.add_run(
    '开发环境搭建包括以下步骤：安装Node.js运行环境（v20 LTS）、使用Midway.js CLI初始化'
    '项目脚手架（midway-init）、安装TypeORM及MySQL驱动（mysql2）、使用Create React App'
    '初始化Web端和管理后台项目、配置微信小程序开发工具。AI工具（Claude Code）通过终端'
    '命令行进行交互，可直接读取项目文件（Read工具）、生成代码文件（Write工具）、'
    '执行Shell命令（Bash工具）和搜索文件内容（Grep/Glob工具）。在开发过程中，首先将'
    '《需求规格说明书》和《视觉设计规范》两个核心文档提供给AI，使其充分理解项目背景和'
    '技术要求，然后按照"数据库设计→后端服务→Web前端→小程序→管理后台→报告撰写"'
    '的顺序逐步推进开发，每完成一个环节进行人工审查后再进入下一个环节。'
)

p = doc.add_paragraph()
p.add_run('3.4.2 数据库层实现')
p = doc.add_paragraph()
p.add_run(
    '数据库层使用TypeORM的装饰器模式定义实体类。以景区实体(ScenicSpot)为例，使用@Entity'
    '装饰器映射到t_scenic_spot表，@PrimaryGeneratedColumn定义自增主键，@Column定义普通字段，'
    '@CreateDateColumn和@UpdateDateColumn自动管理创建和更新时间，@OneToMany定义与票种和评价的'
    '一对多关联关系。实体间的关联关系设计遵循领域驱动设计（DDD）原则，以景区和路线套餐为'
    '聚合根，票种、行程、电子票分别为其子实体。所有实体在entity/index.ts中统一导出，'
    'TypeORM的synchronize选项在开发环境自动同步表结构，生产环境使用Migration管理数据库变更。'
)

p = doc.add_paragraph()
p.add_run('3.4.3 订单服务核心实现')
p = doc.add_paragraph()
p.add_run(
    '订单服务(TravelOrderService)是本模块最核心的业务逻辑。以下为关键业务实现的代码思路：'
    '（1）创建门票订单(createTicketOrder)：查询票种信息（含关联景区）→校验库存→计算总金额'
    '（服务端计算，不信任客户端传入金额）→生成订单号→创建待支付订单→返回订单号和金额。'
    '（2）创建路线订单(createRouteOrder)：查询路线信息→校验提前1天购买规则（R9-03）→'
    '校验库存→计算金额→创建订单。'
    '（3）支付成功回调(paySuccess)：校验订单状态（必须是待支付）→扣减库存（使用UPDATE ...'
    'WHERE stock >= quantity原子操作防止超卖）→更新订单状态为已支付→为每位游客生成电子票'
    '（含唯一二维码、有效期计算）。'
    '（4）退款处理(refund)：校验退款条件（路线出发前3天规则R9-03）→更新状态为退款中→'
    '标记电子票为已退款→由管理员审核（approveRefund）完成退款或驳回。'
    '（5）电子票核销(verifyTicket)：查询电子票→校验状态（未使用/未过期）→校验有效期→'
    '更新为已使用→记录核销人和时间（R9-05）。'
)

p = doc.add_paragraph()
p.add_run('3.4.4 前端页面实现')
p = doc.add_paragraph()
p.add_run(
    'Web前端（PC端）采用React + Ant Design组件库，主要页面包括：'
    '首页（HomePage）：顶部全屏轮播Banner + 搜索框 + 功能入口（景区/路线/攻略/订单）+ '
    '热门景区4列卡片 + 热门路线4列卡片；'
    '景区详情页（ScenicDetailPage）：左栏景区介绍+评价列表，右栏票种选购卡片（含收藏按钮、'
    '购买弹窗、日期选择、数量选择）；'
    '路线详情页（RouteDetailPage）：左栏行程Timeline+详情+注意事项+评价，右栏购买卡片'
    '（含价格、费用包含、库存、预订按钮）；'
    '购买页面（TicketPurchasePage/RoutePurchasePage）：多步骤表单（填写信息→确认订单→支付），'
    '支持动态数量的游客信息录入（R9-04实名制要求）；'
    '订单页面（OrderListPage/OrderDetailPage）：Tab切换订单状态，详情页含电子票二维码展示'
    '（使用Ant Design QRCode组件）、取消和退款操作按钮。'
    '所有页面遵循视觉设计规范的色彩系统（苗银蓝#1F5FA8为主色，苗绣橙#E85D2F为强调色）'
    '和间距系统（8px基准栅格）。'
)

p = doc.add_paragraph()
p.add_run('3.4.5 微信小程序实现')
p = doc.add_paragraph()
p.add_run(
    '微信小程序端采用原生框架开发，共12个页面。使用app.json配置底部TabBar（景区门票/路线套餐/'
    '交通攻略/我的订单四个Tab），app.wxss定义全局样式变量（品牌色、间距、圆角等）。'
    'app.js封装了统一的request方法，自动携带JWT Token和处理错误。'
    '小程序端的关键特性包括：景区/路线列表支持下拉刷新和上拉加载更多；'
    '电子票详情页使用Canvas绘制二维码（R9-02），支持全屏展示便于核销扫码；'
    '购买页面校验提前1天规则（路线）和实名制信息完整性；支付调用微信支付接口（JSAPI）。'
)

p = doc.add_paragraph()
p.add_run('3.4.6 管理后台实现')
p = doc.add_paragraph()
p.add_run(
    '管理后台包含8个功能页面：数据看板（Dashboard，含ECharts柱状图和统计卡片）、'
    '景区管理（ScenicManagePage，CRUD + 上下架开关）、票种管理（TicketManagePage，含景区下拉选择）、'
    '路线套餐管理（RouteManagePage，含行程安排的内联编辑弹窗）、'
    '订单管理（OrderManagePage，含退款审核操作）、电子票核销（VerifyTicketPage，输入票号核销）、'
    '交通攻略管理（GuideManagePage）、评价管理（ReviewManagePage，含商家回复弹窗）。'
    '管理后台采用侧边栏+内容区的经典布局，左侧Sider可折叠，使用Ant Design的Layout、'
    'Menu、Table、Modal、Form等组件构建完整的CRUD操作界面。'
    '数据统计使用ECharts（echarts-for-react）生成柱状图展示热门景区和热门路线排行。'
)

p = doc.add_paragraph()
p.add_run('3.4.7 AI编程经验与软件工程实践')
p = doc.add_paragraph()
p.add_run(
    '在AI辅助编程过程中，本人总结了以下实践经验：'
    '（1）提示词工程（Prompt Engineering）：清晰、结构化的需求描述是高质量AI输出的前提。'
    '在生成数据库实体时，直接引用需求文档中的数据实体定义作为上下文，生成的实体类准确率显著提高。'
    '（2）分层生成策略：将复杂的系统拆解为"实体定义→DTO定义→服务层→控制器→前端页面"'
    '的分层任务，每层独立让AI生成，层与层之间通过人工审查保证接口一致性。'
    '这种方法避免了AI一次性生成过多代码导致上下文混乱的问题。'
    '（3）人工审查不可替代：AI生成的代码在业务规则处理上可能存在遗漏或错误。'
    '例如，路线订单的"提前1天购买"规则，AI最初生成的代码没有正确处理日期比较的边界条件，'
    '经人工审查后发现并修正。因此，AI生成代码后必须进行人工Code Review。'
    '（4）增量迭代优于全量重写：在已有代码基础上通过Edit工具进行局部修改，让AI理解上下文后'
    '进行精确替换，比让AI重新生成整个文件更可靠且更高效。'
    '（5）种子数据的重要性：利用AI生成4个景区、9种票种、4条路线套餐和3条交通攻略的模拟数据，'
    '使得系统在无真实数据的情况下也能进行完整的功能演示和业务流程验证。'
)

# --- 3.5 ---
doc.add_heading('3.5 测试、部署与交付', level=2)

p = doc.add_paragraph()
p.add_run('3.5.1 AI辅助测试用例生成')
p = doc.add_paragraph()
p.add_run(
    '利用AI工具生成了系统化的测试用例，覆盖以下场景：'
    '（1）景区列表接口测试：验证分页参数（page/pageSize）、关键词搜索（keyword）、'
    '排序方式（sortBy: default/price_asc/price_desc/rating），以及景区下架后不显示的规则。'
    '（2）门票下单测试：正常下单流程、票种库存不足（应拒绝下单并提示"库存不足"）、'
    '票种已下架（应返回"票种不存在或已下架"）、游客信息为空（应校验失败）。'
    '（3）路线下单测试：提前1天规则（选择今天或明天出发应拒绝）、库存不足（超过stock数量应拒绝）、'
    '正常下单并生成电子票。'
    '（4）退款测试：路线出发前3天规则（出发前2天申请退款应拒绝）、门票过期后退款应拒绝。'
    '（5）电子票核销测试：正常核销（标记为已使用）、重复核销（应提示"电子票已使用"）、'
    '过期核销（应提示"电子票已过期"）。'
    '通过AI生成的测试用例，发现了库存扣减的并发安全问题（初始使用read-check-write模式），'
    '后修改为数据库原子操作（UPDATE ... WHERE stock >= quantity）。'
)

p = doc.add_paragraph()
p.add_run('3.5.2 部署流程')
p = doc.add_paragraph()
p.add_run(
    '后端服务部署：使用npm run build编译TypeScript为JavaScript，通过node dist/bootstrap.js启动。'
    '配置环境变量：DB_HOST（数据库地址）、DB_PORT（端口）、DB_USER（用户名）、'
    'DB_PASSWORD（密码）、DB_NAME（数据库名）、JWT_SECRET（JWT密钥）、'
    'REDIS_HOST（Redis地址）。生产环境使用PM2进行进程管理和自动重启。'
    '前端应用部署：React项目执行npm run build生成静态文件（build目录），'
    '部署到Nginx的静态文件目录。Nginx配置将/api/路径代理到后端服务，'
    '其他路径返回React应用的index.html（支持BrowserHistory路由）。'
    '微信小程序部署：通过微信开发者工具上传代码→登录微信公众平台→提交审核→审核通过后发布。'
    '管理后台部署方式与Web前端一致，部署在/admin/路径下。'
)

# --- 3.6 ---
doc.add_heading('3.6 模拟运营与业务流程实战', level=2)

p = doc.add_paragraph()
p.add_run('3.6.1 模拟数据生成')
p = doc.add_paragraph()
p.add_run(
    '利用AI工具生成了乌东地区真实的文旅运营数据。种子数据包括：'
    '4个景区：乌东苗寨（苗族木结构建筑群，银饰蜡染非遗体验）、雷公山国家森林公园'
    '（苗岭主峰，云海日出，原始森林）、西江千户苗寨（世界最大苗族聚居村寨，夜景天花板）、'
    '郎德上寨（中国民间艺术之乡，苗族文化活态博物馆）。'
    '9种票种：覆盖4个景区的成人票（¥60-90）、儿童票（¥30）、学生票（¥40-60）、'
    '家庭套票（¥130，2大1小）等。'
    '4条路线套餐：乌东苗寨非遗文化一日游（¥198，含银饰锻造+蜡染体验+长桌宴）、'
    '雷公山自然探索两日游（¥498，含民宿住宿+云海日出）、西江千户苗寨夜景摄影两日游'
    '（¥598，含观景客栈+苗服旅拍）、黔东南苗寨深度研学三日游（¥1280，含三大苗寨+三种非遗体验）。'
    '3条交通攻略：贵阳出发（高铁+大巴，约3h，¥150）、凯里自驾（约1.5h，¥80）、广州出发（高铁，约5.5h，¥380）。'
    '这些数据参考了真实的景区信息和市场定价，确保了模拟运营的真实性和可用性。'
)

p = doc.add_paragraph()
p.add_run('3.6.2 核心业务流程演练（以门票购买为例）')
p = doc.add_paragraph()
p.add_run(
    '完整的门票购买与核销业务流程如下：'
    '步骤1-浏览：游客打开小程序首页→点击"景区门票"Tab→看到乌东苗寨、雷公山等景区列表→'
    '点击"乌东苗寨"进入详情页。'
    '步骤2-选票：在票种列表中看到成人票¥60、儿童票¥30、学生票¥40、家庭套票¥130→'
    '点击成人票的"购买"按钮。'
    '步骤3-确认：弹出购买窗口，显示票种名称（成人票）、单价（¥60）、使用日期选择器→'
    '选择明天（2026-06-28）→数量选择2人。'
    '步骤4-实名：系统展开2组游客信息输入框（R9-04：姓名、身份证号、手机号）→'
    '分别填写"张三 5226xxxxxxxxxxxxxx 138xxxxxxxx"和"李四 5226xxxxxxxxxxxxxx 139xxxxxxxx"→'
    '填写联系人手机号→显示合计金额¥120→点击"确认并支付"。'
    '步骤5-支付：系统调用后端创建订单API→生成订单号WD20260627ABCD1234→返回订单号和金额→'
    '调用微信支付JSAPI→用户输入支付密码→支付成功→微信服务器回调后端→后端验签→'
    '扣减库存（成人票stock由500减为498）→生成2张电子票（票号TKxxx001、TKxxx002，有效期至2026-07-28）。'
    '步骤6-核销：游玩当日，游客在"我的订单"→订单详情→查看电子票二维码→'
    '景区工作人员在管理后台"电子票核销"页面输入票号或扫描二维码→系统验证票状态→'
    '核销成功→票标记为"已使用"→R9-05：核销通知发送给用户。'
)

p = doc.add_paragraph()
p.add_run('3.6.3 模拟运营中暴露的问题与改进方案')
p = doc.add_paragraph()
p.add_run(
    '通过模拟运营，发现以下需要改进的问题：'
    '（1）高并发场景下的库存超卖风险：当前使用数据库行级锁（UPDATE ... WHERE stock >= qty）'
    '来防止超卖，但在极高并发（如节假日抢票）场景下，MySQL的行锁可能成为性能瓶颈。'
    '改进方案：引入Redis分布式锁或使用Redis缓存库存，通过Lua脚本实现原子扣减。'
    '（2）支付回调的幂等性处理：微信支付可能因网络重试发送重复回调，如果未做幂等处理可能导致'
    '重复生成电子票。改进方案：在支付回调中增加事务ID去重检查，使用Redis记录已处理的transaction_id。'
    '（3）电子票过期自动处理：当前电子票过期状态的判断在查询时进行，缺乏主动的定时任务。'
    '改进方案：使用定时任务（如Cron Job）每日凌晨扫描过期电子票并批量更新状态。'
    '（4）退款原路返回：当前退款流程仅更新了订单状态，未对接微信支付退款API。'
    '改进方案：对接微信支付退款接口，实现T+1原路退回。'
    '（5）数据统计实时性：当前数据看板的统计查询（GROUP BY + COUNT）在大数据量时响应较慢。'
    '改进方案：引入定时预计算（物化视图）或使用ClickHouse等OLAP数据库。'
)

# --- 实践总结 ---
doc.add_heading('实践总结与反思', level=2)

p = doc.add_paragraph()
p.add_run(
    '通过本次课程设计，本人对文旅行业的票务业务有了深入而全面的理解。景区门票与路线套餐虽然是'
    '"行"模块的两个子业务，但它们在数据结构、业务规则和用户流程上有显著差异：门票是标准化的商品，'
    '同一票种价格固定，核心关注有效期管理和日库存控制；路线套餐是非标准化的服务商品，包含行程安排、'
    '吃住行等多项内容，核心关注成团人数（库存）和出发日期管理。这两种业务模式的差异要求'
    '在系统设计时采用不同的数据模型和处理逻辑，而非简单的统一抽象。这让我深刻认识到，'
    '在软件工程中，深入理解业务领域并据此选择合适的架构方案，比追求技术上的"统一"更为重要。'
)

p = doc.add_paragraph()
p.add_run(
    '在技术层面，Midway.js框架的依赖注入（IoC）和装饰器模式使得代码组织清晰，模块间耦合度低。'
    'TypeORM的实体管理和Repository模式简化了数据库操作，但在处理复杂查询（如统计聚合）时'
    'QueryBuilder比Repository API更为灵活。React 18的函数组件和Hooks使前端状态管理更加简洁，'
    'Ant Design 5的组件体系大幅提升了UI开发效率。微信小程序的开发受限于其特有的框架和API规范，'
    '但通过WXML的数据绑定和WXSS的样式系统，结合rpx响应式单位，也能实现良好的移动端用户体验。'
    '全栈开发的最大挑战不是精通每一种技术，而是确保各层之间的接口一致性、数据格式统一和'
    '错误处理协调，这需要严格的API文档管理和前后端联调。'
)

p = doc.add_paragraph()
p.add_run(
    '在AI辅助开发方面，本次实践让我深刻体会到AI工具是强大的生产力倍增器，但它不能替代'
    '开发者的工程判断力。AI在代码生成、重复性任务、文档撰写等方面效率极高，但在需求理解、'
    '架构设计、安全风控、业务规则校验等需要深度思考的环节，人类的专业判断仍然不可替代。'
    '真正的"AI全栈开发"不是让AI替代人，而是人利用AI更高效地完成工作，将精力聚焦在更有创造性的问题上。'
    '通过本次实践，本人不仅掌握了Midway.js+React+微信小程序的全栈开发技能，'
    '更重要的是培养了利用AI工具进行高效开发的思维方式和工程实践能力，'
    '这将对未来的学习和职业发展产生深远影响。'
)

p = doc.add_paragraph()
p.add_run(
    '本项目的不足与未来展望：（1）当前未实现真实的微信支付对接，支付环节使用了模拟回调，'
    '后续需要接入微信支付商户API，完成完整的支付-回调-退款链路；（2）未实现短信通知和'
    '微信订阅消息推送功能，用户体验有待提升；（3）缺少图片上传和OSS存储的完整对接，'
    '当前使用占位URL替代真实图片；（4）数据统计看板仅展示了基础指标，可进一步丰富分析维度'
    '（如用户转化漏斗、时段分布、地域分布等）；（5）未进行大规模并发测试，系统的性能瓶颈'
    '和容量上限尚不明确，建议使用JMeter进行压测；（6）前端缺少骨架屏加载和更丰富的交互动效，'
    '用户体验可进一步优化。未来还可以考虑引入AI推荐算法，根据用户的浏览历史、收藏偏好和出行季节，'
    '智能推荐最适合的景区和路线套餐，实现个性化旅游规划。'
)

# --- 参考文献 ---
doc.add_heading('参考文献', level=2)

refs = [
    '[1] Midway.js团队. Midway.js官方文档 v3.x[EB/OL]. https://midwayjs.org, 2024.',
    '[2] React官方团队. React官方文档 v18.x[EB/OL]. https://react.dev, 2024.',
    '[3] 腾讯微信团队. 微信小程序开发文档[EB/OL]. https://developers.weixin.qq.com/miniprogram/dev/, 2024.',
    '[4] 蚂蚁集团. Ant Design 5.x 组件库文档[EB/OL]. https://ant.design, 2024.',
    '[5] TypeORM团队. TypeORM官方文档[EB/OL]. https://typeorm.io, 2024.',
    '[6] 王佛琴. 人工智能工程实践课程教学大纲[Z]. 厦门理工学院计算机工程学院, 2026.',
    '[7] 乌东文旅项目组. 乌东文旅衣食住行需求规格说明书V1.0[Z]. 2026.',
    '[8] 乌东文旅项目组. 乌东文旅衣食住行视觉设计规范V1.0[Z]. 2026.',
]

for ref in refs:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(0)
    p.add_run(ref)

# 保存
output_path = '人工智能工程实践-课程设计大作业报告_已完成.docx'
doc.save(output_path)

# 统计字数
total_chars = sum(len(p.text) for p in doc.paragraphs)
print(f'Report saved to: {output_path}')
print(f'Total characters: {total_chars}')
print('Done!')
