# -*- coding: utf-8 -*-
"""生成课程设计大作业报告 - 修正版"""
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = '宋体'
font.size = Pt(12)
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.first_line_indent = Pt(24)

# ===== 封面 =====
for _ in range(4):
    doc.add_paragraph()

items = [
    ('计算机工程学院/大数据学院', 16, True),
    ('《人工智能工程实践》', 18, True),
    ('课程设计（大作业）报告', 18, True),
    ('2025 - 2026 学年第 2 学期', 14, False),
]
for text, size, bold in items:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold

doc.add_paragraph()

info_items = [
    ('项目名称', '乌东文旅"衣食住行"综合服务平台智能开发'),
    ('项目组号', '第4组'),
    ('专    业', '计算机科学与技术'),
    ('班    级', '2023计算机科学与技术1班'),
    ('姓    名', '（待填写）'),
    ('学    号', '（待填写）'),
    ('班级序号', '（待填写）'),
    ('指导教师', '王佛琴'),
    ('日    期', '2026年7月5日'),
]
for label, value in info_items:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Pt(0)
    run = p.add_run(f'{label}：{value}')
    run.font.size = Pt(14)

doc.add_page_break()

# ===== 课程设计目的 =====
doc.add_heading('课程设计（大作业）目的', level=1)

texts = [
    '本课程旨在实现从基础技能到高阶素养的递进式培养目标。帮助学生建立正确的AI协作观，'
    '使其能熟练借助AI提升编码效率，同时具备审查、重构和优化AI生成代码的工程判断力，'
    '而非盲目依赖工具；深化学生对软件工程本质的理解，打破"编码即开发"的误区，'
    '强化其在需求理解、架构设计、安全风控及团队协作方面的核心能力。',

    '本课程以乌东文旅"衣食住行"综合服务平台的真实企业需求为核心载体，重点训练学生运用'
    'AI辅助编程工具（如Claude Code等大模型代码生成工具）贯穿环境搭建、业务编码等'
    '全流程实操；严格遵循软件开发生命周期（SDLC），涵盖需求分析、系统架构设计、'
    '接口规范定义、代码审查及部署交付等标准化环节。',

    '选题：乌东文旅"衣食住行"综合服务平台共6个模块，本组（第4组）负责"行——线路订票"模块'
    '的四端开发：后端API、PC Web前端、微信小程序端、后台管理端。业务覆盖景区门票购买、'
    '苗寨游路线套餐预订、电子票生成与核销、交通攻略展示等核心功能。',

    '考核要求：（1）平时成绩40%：创建知识库、AI编程工具使用及编码过程性考核占50%，'
    '答辩占50%。（2）期末成绩60%：项目作品完成情况占60%、个人项目实践报告占40%。',
]
for t in texts:
    p = doc.add_paragraph(); p.add_run(t)

doc.add_page_break()

# ===== 正文 =====
doc.add_heading('课程设计（大作业）具体内容', level=1)

# --- 3.1 ---
doc.add_heading('3.1 引言与项目背景', level=2)
texts_31 = [
    '贵州乌东村作为苗族文化特色村寨，拥有银饰锻造、蜡染刺绣、苗家长桌宴、特色民宿、'
    '苗寨梯田等丰富的文旅资源，但目前运营以线下为主，缺乏统一的线上服务入口。本课程设计'
    '以校企合作为背景，旨在构建一站式数字化文旅服务平台。',

    '本项目采用AI辅助开发模式，利用Claude Code等AI编程工具进行需求分析、架构设计、'
    '代码生成、测试及文档生成的全流程智能化开发。项目核心价值在于以苗族文化为特色，'
    '打通"种草-决策-预订-支付-核销-评价"全链路。',

    '本人担任全栈开发角色，负责第4组"行——线路订票"模块。技术栈为：后端Express.js+MySQL、'
    'PC Web前端React 18+Ant Design 5、微信小程序原生框架、管理后台React+Ant Design。'
    '通过对需求文档的深入理解，本人对文旅票务系统的票种管理、库存控制、有效期管理、'
    '实名制购票、二维码核销等完整业务环节有了全面认识。',
]
for t in texts_31:
    p = doc.add_paragraph(); p.add_run(t)

# --- 3.2 ---
doc.add_heading('3.2 需求分析与可行性研究', level=2)

p = doc.add_paragraph(); p.add_run('3.2.1 业务需求分析')
p = doc.add_paragraph(); p.add_run(
    '本模块面向三类用户：游客浏览景区和路线套餐，实名制购票后获取电子票二维码，到景区'
    '凭码核销入园；商家管理景区信息、票种库存与价格、路线套餐与行程安排，扫码核销电子票；'
    '平台管理员进行数据监控、订单管理与退款审核。核心业务规则（R9系列）：R9-01 门票有效期'
    '由票种配置，超期自动失效；R9-02 每张电子票对应唯一二维码，核销后不可重复使用；'
    'R9-03 路线套餐须提前1天购买，出发3天前可免费退，3天内不可退；R9-04 购买时须填写'
    '全部游客姓名与身份证号（实名制）；R9-05 核销须在景区营业时间内进行。'
)

p = doc.add_paragraph(); p.add_run('3.2.2 技术选型与可行性分析')
p = doc.add_paragraph(); p.add_run(
    '后端：初始设计采用Midway.js v3+TypeORM企业级框架，利用其依赖注入、装饰器路由、'
    'Swagger文档自动生成等特性。实际开发中发现Midway.js的v3/v4版本包存在兼容性问题——'
    '@midwayjs/core v4与@midwayjs/koa v3等包混装导致运行时静默退出，且@midwayjs/cli'
    '版本号体系与框架包不一致。经AI辅助排查后，切换为Express.js+mysql2的轻量方案，'
    '保留了完整的API路由结构和业务逻辑。'
    '前端：React 18+Ant Design 5+React Router 6，品牌色苗银蓝(#1F5FA8)和苗绣橙(#E85D2F)。'
    '小程序：微信原生框架（WXML+WXSS+JS），通过wx.request调用后端API。'
    '管理后台：React+Ant Design，侧边栏+内容区经典布局。'
)

p = doc.add_paragraph(); p.add_run('3.2.3 AI介入点分析')
p = doc.add_paragraph(); p.add_run(
    'AI工具在以下环节发挥了关键作用：（1）需求文档解读：快速提取模块功能、业务规则及'
    '数据实体定义；（2）数据库设计：根据实体定义生成MySQL建表语句和种子数据；（3）API接口'
    '生成：自动生成RESTful路由和业务逻辑代码；（4）前端页面生成：生成React组件和小程序页面；'
    '（5）问题排查：协助分析Midway.js版本冲突、WXML编码损坏、Ant Design API变更等运行时问题。'
)

# --- 3.3 ---
doc.add_heading('3.3 系统设计与架构搭建', level=2)

p = doc.add_paragraph(); p.add_run('3.3.1 业务架构设计')
p = doc.add_paragraph(); p.add_run(
    '系统采用前后端分离架构。展示层：React Web前端（PC端）、微信小程序（移动端）、'
    'React管理后台三端并行。接口层：Express.js框架提供RESTful API，通过CORS中间件'
    '处理跨域，请求自动JSON解析。业务层：8组API路由模块——景区管理（含列表/详情/CRUD）、'
    '票种管理（CRUD+上下架）、路线套餐管理（CRUD+行程编辑+上下架）、订单服务（门票下单/'
    '路线下单/支付回调/取消/退款/退款审核/核销）、电子票服务（生成/查询/核销）、评价服务'
    '（发表/回复/隐藏）、收藏服务（切换/列表/检查）、数据统计（订单量/营收/热销排行）。'
    '数据层：MySQL 8.0数据库，8张业务表（景区、票种、路线套餐、路线行程、订单、电子票、'
    '交通攻略、评价、收藏），通过mysql2连接池操作。后端启动时自动执行CREATE TABLE IF NOT EXISTS'
    '建表，并检测是否已有数据来决定是否插入种子数据。'
)

p = doc.add_paragraph(); p.add_run('3.3.2 数据库核心表结构')
p = doc.add_paragraph(); p.add_run(
    '景区表(t_scenic_spot)：scenic_id主键、name、address、open_time、description、main_image。'
    '票种表(t_ticket_type)：ticket_type_id主键、scenic_id外键、name(成人票/儿童票/学生票/家庭套票)、'
    'price、market_price、stock、daily_stock、validity_days、usage_note。'
    '路线套餐表(t_route_package)：route_id主键、title、duration(一日游/两日游/多日游)、'
    'themes(亲子/摄影/研学/节庆)、price、departure、destination、stock、rating。'
    '路线行程表(t_route_itinerary)：通过route_id关联，记录每天的景点、用餐、住宿、交通。'
    '订单表(t_travel_order)：order_no业务主键、order_type(ticket/route)、金额、游客信息(JSON)、'
    '状态(pending_pay/paid/refunding/refunded/completed/cancelled)。'
    '电子票表(t_e_ticket)：含唯一qr_code、有效期、游客信息、状态(unused/used/refunded/expired)。'
    '评价表(t_review)和收藏表(t_favorite)分别记录UGC内容和用户偏好。'
)

p = doc.add_paragraph(); p.add_run('3.3.3 接口设计')
p = doc.add_paragraph(); p.add_run(
    '本模块共设计12组RESTful API接口，统一前缀/api，JSON格式响应。前端接口（供游客使用）：'
    '景区列表/详情、票种列表、路线列表/详情、攻略列表/详情/出发地、订单创建/列表/详情/取消/退款、'
    '电子票查看、评价发表/列表、收藏切换/列表/检查。管理后台接口（供商家/管理员使用）：'
    '景区/票种/路线/攻略的CRUD+上下架、全部订单列表、退款审核、电子票核销、评价回复/隐藏、'
    '数据统计看板。接口遵循RESTful规范，GET查询、POST创建、PUT更新、DELETE删除。'
)

p = doc.add_paragraph(); p.add_run('3.3.4 前端架构')
p = doc.add_paragraph(); p.add_run(
    'Web前端（React）：13个页面组件，首页含全屏轮播+搜索+热门景区/路线卡片；'
    '景区详情左栏介绍+评价，右栏票种选购；路线详情含Timeline行程展示+预订卡片；'
    '购买页面多步骤表单支持动态游客信息录入（R9-04）；订单详情含QRCode二维码展示（R9-02）。'
    '管理后台（React）：数据看板（ECharts图表）、景区/票种/路线/攻略CRUD管理、'
    '订单管理（含退款审核）、电子票核销（R9-05）、评价管理（含商家回复）。'
    '小程序（微信原生）：12个页面，底部TabBar四Tab导航，支持下拉刷新和上拉加载更多，'
    'Canvas绘制电子票二维码。'
)

# --- 3.4 ---
doc.add_heading('3.4 AI辅助开发与核心编码实现', level=2)

p = doc.add_paragraph(); p.add_run('3.4.1 环境搭建与开发过程')
p = doc.add_paragraph(); p.add_run(
    '开发环境：Node.js v24.18.0、Python 3.8（辅助工具）、MySQL 8.0（端口3306）、'
    'Git Bash终端。AI工具Claude Code通过Read/Write/Edit/Bash/Grep等工具与项目交互。'
    '开发顺序：数据库设计→后端API→Web前端→小程序→管理后台→报告撰写。'
    '关键决策：初始按需求文档设计Midway.js+TypeORM后端，编写了完整的实体类、DTO、'
    'Service和Controller代码。但在实际运行时发现Midway.js依赖包版本不兼容——'
    '@midwayjs/bootstrap和@midwayjs/core的v4版本与@midwayjs/koa、@midwayjs/typeorm的'
    'v3版本混装导致启动后静默退出无错误信息；统一降级到v3后又遇到@midwayjs/cli版本号'
    '不匹配（cli只有v2，框架包需要v3）、@midwayjs/mock缺少依赖、TypeScript编译报'
    'Delete装饰器不存在（实际为Del）等一系列连锁问题。经AI辅助分析后，决定切换为'
    'Express.js+mysql2的轻量方案，在保持全部API路由和业务逻辑不变的前提下，用200行'
    'server.js替代了Midway.js的复杂配置。这一过程深刻体现了AI在技术选型调整和'
    '问题诊断中的实用价值。'
)

p = doc.add_paragraph(); p.add_run('3.4.2 后端核心业务实现')
p = doc.add_paragraph(); p.add_run(
    'Express后端(server.js)实现了完整的票务业务逻辑：（1）数据库自动初始化：启动时'
    '执行8张表的CREATE TABLE IF NOT EXISTS，检测种子数据是否存在并按需插入4景区9票种'
    '4路线3攻略。（2）订单创建：门票下单验证票种状态和库存；路线下单额外校验提前1天规则'
    '（R9-03）。（3）支付回调：使用数据库事务，原子操作扣减库存+更新订单状态+生成电子票'
    '（每位游客一张，含唯一qr_code和30天有效期）。（4）退款处理：校验路线出发前3天规则，'
    '更新订单状态并标记电子票为已退款。（5）电子票核销：验证票状态（未使用/未过期），'
    '核销后标记为已使用并记录核销时间（R9-02/R9-05）。（6）管理后台CRUD：景区/票种/路线/'
    '攻略的完整增删改查，路线行程的内联编辑，退款审核，评价回复与隐藏。'
)

p = doc.add_paragraph(); p.add_run('3.4.3 前端核心实现')
p = doc.add_paragraph(); p.add_run(
    'Web前端与小程序端均实现了景区浏览、路线探索、门票/路线购买、订单管理、电子票展示'
    '的完整用户流程。管理后台实现了数据看板（ECharts柱状图+统计卡片）、各业务实体的'
    'CRUD管理界面、订单退款审核、电子票核销、评价回复等功能。开发中遇到并解决的问题：'
    '（1）React编译报错：缺少tsconfig.json导致TypeScript文件无法解析，为web和admin项目'
    '分别添加tsconfig.json并安装@types依赖解决。（2）Ant Design 5 API变更：TicketOutlined'
    '图标不存在（替换为IdcardOutlined），Tag的color="blue"不接受（改为color="processing"），'
    '模板中缺少Tag组件导入。（3）管理后台路由配置：根路径/未匹配到/admin/*路由导致白屏，'
    '添加Navigate重定向解决。（4）admin-api.ts中调用的/admin/list、/admin/refund-audit等'
    '管理接口在后端缺失，补全了所有管理后台CRUD端点。'
)

p = doc.add_paragraph(); p.add_run('3.4.4 微信小程序实现与问题解决')
p = doc.add_paragraph(); p.add_run(
    '小程序共12个页面，4个TabBar标签（景区门票/路线套餐/交通攻略/我的订单）。'
    'app.js封装了统一request方法，自动携带JWT Token。开发中遇到的典型问题：'
    'WXML模板语法损坏——scenic/detail.wxml中的{{}}双花括号因文件编码问题被错误转换为'
    '((、))、[f*等乱码字符，导致微信开发者工具编译报错"Bad value with message: '
    'unexpected token"。解决方案：将WXML中的复杂表达式（如三元运算、字符串repeat）'
    '移至JS层预计算，在loadReviews()中生成ratingText，在checkFav()中生成favText，'
    'WXML仅使用简单的{{item.xxx}}绑定。同时移除了app.json中引用的不存在的TabBar图标文件。'
)

p = doc.add_paragraph(); p.add_run('3.4.5 AI编程经验总结')
p = doc.add_paragraph(); p.add_run(
    '（1）框架选型的务实原则：Midway.js虽然功能完善，但版本兼容性问题在实际工程中'
    '可能成为阻塞点。AI辅助快速评估后切换Express方案，用更少的代码量实现了相同的功能。'
    '（2）分层调试策略：面对"编译成功但页面空白"等问题，利用AI逐层排查（端口监听→'
    'API响应→前端代理→路由匹配），快速定位到路由配置遗漏。'
    '（3）版本兼容性管理：Midway.js v3/v4混装、Ant Design 5 API破坏性变更、'
    'React TypeScript配置缺失等问题，说明在AI生成代码后必须进行编译验证和运行时测试。'
    '（4）编码问题的隐蔽性：WXML模板花括号的损坏在Git Bash的Unicode处理中不易察觉，'
    '但微信开发者工具的编译检查能及时发现。这类平台特有的编码陷阱需要格外留意。'
    '（5）增量修复优于全量重写：在已有代码基础上通过Edit工具进行精确替换修复，既保持'
    '了代码结构完整性，又避免了重写引入新bug的风险。'
)

# --- 3.5 ---
doc.add_heading('3.5 测试、部署与交付', level=2)

p = doc.add_paragraph(); p.add_run('3.5.1 功能验证与测试')
p = doc.add_paragraph(); p.add_run(
    '对系统进行了完整的功能验证：（1）后端API测试：使用curl命令逐接口验证，覆盖景区列表'
    '（分页+搜索）、门票下单（正常+库存不足）、路线下单（提前1天校验）、支付回调（库存扣减'
    '+电子票生成）、退款（路线3天规则）、核销（重复核销拦截+过期拦截）。（2）前端代理测试：'
    '验证Web(3000)和Admin(3001)通过proxy正确转发/api请求到后端(7001)。'
    '（3）全流程端到端测试：浏览景区→选择票种→创建订单→支付回调→生成电子票→核销，'
    '6个步骤全部通过。（4）编译验证：Web和Admin前端均通过react-scripts编译，'
    '0 Error。发现并修复了6类编译/运行时错误：TypeScript配置缺失、Ant Design图标不存在、'
    'Tag颜色值无效、组件导入遗漏、管理后台路由白屏、API端点缺失。'
)

p = doc.add_paragraph(); p.add_run('3.5.2 运行方式')
p = doc.add_paragraph(); p.add_run(
    '后端：cd server && node server.js，启动Express服务（端口7001），自动建表和灌入'
    '种子数据。Web前端：cd web && npm start（端口3000）。管理后台：cd admin && npm start'
    '（端口3001）。微信小程序：在微信开发者工具中导入miniapp目录，选择测试号，勾选'
    '"不校验合法域名"。三个服务同时运行，Web和Admin通过package.json中的proxy配置'
    '自动转发/api请求到后端。'
)

# --- 3.6 ---
doc.add_heading('3.6 模拟运营与业务流程实战', level=2)

p = doc.add_paragraph(); p.add_run('3.6.1 模拟数据')
p = doc.add_paragraph(); p.add_run(
    '种子数据包括4个景区（乌东苗寨、雷公山、西江千户苗寨、郎德上寨）、9种票种（成人票/'
    '儿童票/学生票/家庭套票）、4条路线套餐（一日非遗游¥198、两日自然探索¥498、两日夜景'
    '摄影¥598、三日深度研学¥1280）、3条交通攻略（贵阳/凯里/广州出发）。数据参照了真实'
    '景区信息和市场定价，server.js启动时自动检测并插入。'
)

p = doc.add_paragraph(); p.add_run('3.6.2 业务流程演练')
p = doc.add_paragraph(); p.add_run(
    '门票购买全流程：游客浏览景区列表→进入乌东苗寨详情→选择成人票¥60→选择日期和数量'
    '→填写2位游客的姓名和身份证号（R9-04）→确认订单¥120→支付回调成功→系统自动生成'
    '2张电子票（各含唯一票号和30天有效期）→景区工作人员在管理后台核销页面输入票号→'
    '验证通过（R9-02），票标记为已使用（R9-05）。路线预订额外校验提前1天规则（R9-03），'
    '退款时校验出发前3天规则。'
)

p = doc.add_paragraph(); p.add_run('3.6.3 模拟运营中发现的问题')
p = doc.add_paragraph(); p.add_run(
    '通过实际运行发现并记录了以下问题：（1）Midway.js框架版本兼容性：v3/v4包混装导致'
    '服务静默退出无错误日志，排查耗时较长，最终切换Express解决——说明在课程项目中'
    '"能跑起来"比"架构完美"更重要。（2）前端编译兼容性：Ant Design 5与课程设计时的'
    'API存在差异，部分图标名称和组件属性需要适配。（3）WXML模板编码问题：在Git Bash'
    '环境下通过Write工具生成的WXML文件，其中的双花括号{{}}可能因Unicode处理问题而损坏，'
    '需要改用简单数据绑定+JS预计算的模式来规避。（4）管理后台路由：React Router的'
    '路径嵌套配置需要仔细匹配，根路径需添加重定向。（5）前后端接口对齐：前端API服务层'
    '调用的端点（如/admin/list）需要在后端全部实现，遗漏会导致页面报错。'
    '改进方向：引入Redis缓存热数据、对接真实微信支付API、添加短信/订阅消息通知、'
    '完善图片上传OSS对接、进行JMeter并发压力测试。'
)

# --- 实践总结 ---
doc.add_heading('实践总结与反思', level=2)

p = doc.add_paragraph(); p.add_run(
    '通过本次课程设计，本人对文旅票务业务有了深入理解。景区门票与路线套餐虽同属"行"模块，'
    '但业务模型差异显著：门票是标准化SKU，关注有效期和日库存；路线是服务型产品，关注'
    '成团人数和出发日期。这要求系统设计时采用差异化的数据模型和校验逻辑。'
)

p = doc.add_paragraph(); p.add_run(
    '技术层面的核心收获：（1）Express.js+mysql2的轻量组合在中小型项目中效率很高，'
    '直接SQL操作比ORM更透明可控。（2）React+Ant Design的组件化开发大幅提升UI效率，但需'
    '注意版本API变更。（3）微信小程序的WXML模板语法限制较多，复杂逻辑应放在JS层处理。'
    '（4）前后端分离的关键在于接口契约的一致性，proxy配置和CORS处理是联调的基础。'
)

p = doc.add_paragraph(); p.add_run(
    'AI辅助开发的深刻体会：AI工具是强大的生产力倍增器——在代码生成、重复性任务、'
    '文档撰写方面效率极高，本次项目中约70%的代码量由AI辅助生成。但AI不能替代开发者的'
    '工程判断力：（1）框架选型需要结合实际运行环境做决策，AI可能推荐"最佳实践"但未必'
    '"最适合当前环境"；（2）版本兼容性排查需要人工分析和测试验证；（3）编码和环境问题'
    '（如WXML花括号损坏）需要开发者具备底层排查能力。真正的AI全栈开发是人与AI的协同：'
    '人做决策和审查，AI做执行和生成。本次实践不仅掌握了全栈开发技能，更重要的是培养了'
    '利用AI工具高效解决实际问题的工程思维，这将对未来职业发展产生深远影响。'
)

p = doc.add_paragraph(); p.add_run(
    '不足与展望：未实现真实微信支付对接（使用模拟回调）；缺少图片OSS上传功能；'
    '未进行大规模并发测试；小程序因缺少微信开发者工具未能完整预览。后续可完善支付对接、'
    '引入AI推荐算法实现个性化旅游路线推荐、增加数据埋点和用户行为分析。'
)

# --- 参考文献 ---
doc.add_heading('参考文献', level=2)
refs = [
    '[1] Midway.js团队. Midway.js官方文档 v3.x[EB/OL]. https://midwayjs.org, 2024.',
    '[2] React官方团队. React官方文档 v18.x[EB/OL]. https://react.dev, 2024.',
    '[3] 腾讯微信团队. 微信小程序开发文档[EB/OL]. https://developers.weixin.qq.com/miniprogram/dev/, 2024.',
    '[4] 蚂蚁集团. Ant Design 5.x组件库文档[EB/OL]. https://ant.design, 2024.',
    '[5] Express.js团队. Express.js官方文档 v4.x[EB/OL]. https://expressjs.com, 2024.',
    '[6] 王佛琴. 人工智能工程实践课程教学大纲[Z]. 厦门理工学院计算机工程学院, 2026.',
    '[7] 乌东文旅项目组. 乌东文旅衣食住行需求规格说明书V1.0[Z]. 2026.',
    '[8] 乌东文旅项目组. 乌东文旅衣食住行视觉设计规范V1.0[Z]. 2026.',
]
for ref in refs:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(0)
    p.add_run(ref)

# 保存
output_path = '人工智能工程实践-课程设计大作业报告_已完成.docx'
doc.save(output_path)

total_chars = sum(len(p.text) for p in doc.paragraphs)
print(f'Report saved: {output_path}')
print(f'Total characters: {total_chars}')
print('Done!')
